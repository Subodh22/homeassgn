from docx import Document
import openpyxl
from random import randint
book = openpyxl.load_workbook('C:/Users/Subodh Maharjan/Desktop/fus/Rec.xlsx')

sheet = book.active
import os 
doc= Document('logic.docx')
table1 = doc.tables[0]
table2 = doc.tables[1]
table3 = doc.tables[2]
table4 = doc.tables[3]
table5=doc.tables[5]
table6=doc.tables[6]
table7=doc.tables[7]
table8=doc.tables[8]
table9=doc.tables[9]
table10=doc.tables[10]
table11=doc.tables[11]
time=8.30
supplier=["Bill Veggies and More","Melbourne Marketplace","Green Grocers","George Wholesellers","S and M brothers"]
tool=["Casserole","Oven","Microwave","Pan","Flat Grill","Frier","Wooden Spoons","Measuring cup","Scale"]
food_bank=[]

def tablee():
    
    for x in range(6): 
        num_val = randint(1,9)
        if num_val not in food_bank : 
            food_bank.append(num_val)
        else: 
             continue
    print(food_bank)
    for z in range(len(food_bank)):
       
        a=food_bank[z]
        if(food_bank[z]%2==0):
            a=food_bank[z]-1
        food_name = sheet.cell(row=2,column=a)
        
        y=z+2
        rando_protion=randint(10,30)
        if(y>len(food_bank)):
            continue
        table1.cell(y,1).text=food_name.value
        table1.cell(y,2).text=str(rando_protion)
    tab_instruct=[]
    for b in range(15):
        instr_rando=randint(1,4)
     
        food_bank_val=food_bank[instr_rando]
        if(food_bank_val%2!=0):
            food_bank_val=food_bank_val+1
            q=4
        if food_bank_val not in tab_instruct:
            q=4
            
        else:
            
            q=tab_instruct.count(food_bank_val)+5
        tab_instruct.append(food_bank_val)

        food_instruct=sheet.cell(row=q,column=food_bank_val)
        print(food_instruct.value)
        print(q)
        if food_instruct.value is None :
            continue
        g=b+9
        print(len(table1.rows))
        print(g)
        if(g==len(table1.rows)):
            continue
        table1.cell(g,1).text= food_instruct.value

        end_of_service=["We turned off all the gas stoves and friers","We started cleaning down and wiping all the benches and stoves","We started to sweeping and brooming"]
        criteria=["The criteria for selecting the best ingredients are freshness, temprature, quality, color  "]
        storage=["Before storing any dishes, its temprature is check and recorded and all the containers are labeled and dated"]
        quality_checks=["For quality of all the dishes prepared are tasted and checked by the headchef."]       
        hygiene =["Before preparation of any dishes all chefs must wash their hands and wear proper gloves and ppe"]
        table2.cell(6,1).text=end_of_service[0]
        table2.cell(7,1).text=end_of_service[1]
        table2.cell(8,1).text=end_of_service[2]
        table2.cell(10,1).text=criteria[0]
        table2.cell(15,1).text=hygiene[0]
        table2.cell(19,1).text=quality_checks[0]
        table2.cell(23,1).text=storage[0]
        
       


def table_mise():
    
    tools=[]
    tool_name=""
    for i in range(8):
        randtools=randint(0,6)
        if(randtools in tools):
            continue
        else:
            tools.append(randtools)
            tool_name=tool_name+"\n"+tool[randtools]
    table3.cell(0,0).tables[0].cell(1,0).text=tool_name
    rand_recipe=randint(1,9)
    rec_name=""
    a=rand_recipe
    b=rand_recipe
    if(rand_recipe%2==0):
        a=rand_recipe-1
        rec_name = sheet.cell(row=2,column=a)
    else: 
        b=b+1

    recipe_list=""
    recipe_ins=""
    for i in range(10):
        x=i+4
        cone=sheet.cell(row=x,column=a)
        if cone.value is None :
         continue
        else:
         recipe_list=recipe_list+"\n"+str(cone.value)
        
    table3.cell(0,1).tables[0].cell(1,0).text=recipe_list
    
    for i in range(5):
        x=i+4
        conee=sheet.cell(row=x,column=b)
        if cone.value is None :
         continue
        else:
            if(i==0):
                recipe_ins=str(conee.value)
            else:
                recipe_ins=recipe_ins+str(conee.value)
    
    table4.cell(0,1).text=recipe_ins
        



def order_list():

    list_order=[]
    for i in range(9):
        rand_item=randint(1,9)
        rand_row=randint(4,13)
        if(rand_item%2==0):
            rand_item=rand_item-1
        
        item_name=sheet.cell(row=rand_row,column=rand_item)
        print (item_name.value)
        if item_name.value is None :
             item_name=sheet.cell(row=randint(4,8),column=rand_item)
        item_jus_name=str(item_name.value).split(",")

        if item_jus_name[1] in list_order : 
            continue
        else:
            k=i+1
            list_order.append(item_jus_name[1])
            table6.cell(k,0).text=item_jus_name[1]
            rand_quantity=randint(20,50)
            table6.cell(k,1).text=str(rand_quantity)
            rand_suplier=randint(0,4)
            table6.cell(k,2).text=supplier[rand_suplier]
            
def central_order():
    table7.cell(0,1).text="erwer"
    list_order=[]
    for i in range(10):
        rand_item=randint(1,9)
        rand_row=randint(4,13)
        if(rand_item%2==0):
            rand_item=rand_item-1
        
        item_name=sheet.cell(row=rand_row,column=rand_item)
        print (item_name.value)
        if item_name.value is None :
             item_name=sheet.cell(row=randint(4,8),column=rand_item)
        item_jus_name=str(item_name.value).split(",")

        if item_jus_name[1] in list_order : 
            continue
        else:
            k=i+2
            list_order.append(item_jus_name[1])
            table7.cell(k,0).text=item_jus_name[1]
            rand_quan=randint(0,3)
            Rand_variant=randint(2,10)
            rand_cent=randint(15,40)
            rand_act=rand_cent-Rand_variant
            if(rand_quan == 0):
                rand_act=rand_cent

            Rand_variant=rand_cent-rand_act
            table7.cell(k,1).text=str(rand_cent)
            table7.cell(k,2).text=str(rand_act)
            table7.cell(k,3).text=str(Rand_variant)
            Rand_stock=randint(20,30)
            table7.cell(k,4).text=str(Rand_stock)
            rand_diff=rand_act-Rand_stock
            if(rand_diff>0):
                table7.cell(k,5).text="Update Central Record"
                if(Rand_variant==0):
                    table7.cell(k,5).text="No action Required"
               
            else:

                table7.cell(k,5).text="Place order"


def temp_record():
   
   for i in range(9):
       if(i==0):
           continue
       l=2*i-1
       
       print(l)
       food_item=sheet.cell(row=4,column=l)
       print(food_item.value)
       foodie=str(food_item.value).split(",")
       print(foodie[1])
       k=i+1
       table8.cell(k,0).text=foodie[1]
       rand_tempe=randint(0,3)
       table8.cell(k,1).text=str(rand_tempe)+"C"

       
       
        



def food_prep_sheet():
  prep_item=["Brown Stock, 5G,10G","White Stock,10G,15G","Filets ,20,40","Wild Rice Strudel,10,20","Scallop Potato Pies,25,60","Minced Garlic,-,1/6pan","Minced Shallots,-,1/6 pan","Lobster Stock,-,10 gal","Beef Demi,-,2 gal","Fumet,-,3 gal","Lemon Aioli,1 qt,2 qt","Ceaser Dressing,-,1 gal","Oriental Mignonette,-,1 gal","Joannes Tarter,-,1 gal","Clam Fritters,-,3 1/6 pan","Crab Cakes,18,35","Cajun Tarter,-,1 gal","Bacon-Red Wine Vinegrette,-,1 gal"]
  
  item_rec=[]
  for i in range(20):
      rand_prep_item=randint(0,len(prep_item)-1)
      o=i+3
      if rand_prep_item in item_rec : 
            continue
      else:
          item_rec.append(rand_prep_item)
          if prep_item[rand_prep_item] is None : 
              continue
          item_prep_name=prep_item[rand_prep_item].split(",")
          
          table5.cell(o,0).text=item_prep_name[0]
          table5.cell(o,3).text=item_prep_name[2]
          table5.cell(o,2).text=item_prep_name[1]

def Writer(cellno):
    rec_timer=0
    p=cellno+1
    if(p%3==0):
            rec_timer=0
    rec_timer=+1
    fri=rec_timer+6

    
    
    print(len(table9.rows))
    People=["SM","PO","MM","GR","BM"]
    People_full=["Subodh Maharjan","Param Oli","Mukesh Magar","Grin Ranger","Bishal Muli"]
   
   
    for i in range(len(table9.rows)):
        ol=i+3
        rand_person=randint(0,4)
        if(ol>=len(table9.rows)):
            continue
        print(i,ol,rand_person,len(table9.rows),People[rand_person])
        table9.cell(ol,fri).text=People[rand_person]

    for i in range(len(table10.rows)):
     
        table10.cell(i,fri).text=People[rand_person]
    for i in range(len(table11.rows)):
        k=i+1
        rand_people_full=randint(0,4)
        if(k>=len(table11.rows)):
            continue
        table11.cell(k,3).text=People_full[rand_people_full]
    



def flame_bitches():
    
    for i in range(12):

        tablee()
        table_mise()
        order_list()
        central_order()
        temp_record()
        food_prep_sheet()
        Writer(i)
        pl=str(i)
        doc.save('file'+pl+'.docx')
      

        

        
flame_bitches()

